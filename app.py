# import os
# import fitz
# import pytesseract
# from docx import Document
# from PIL import Image
# from sentence_transformers import SentenceTransformer
# import numpy as np
# import pickle
# import faiss

# pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

# EMBED_MODEL = SentenceTransformer("all-MiniLM-L6-v2")



# def parse_document(file_path):
#     ext = os.path.splitext(file_path)[1].lower()
#     chunks = []

#     if ext == ".pdf":
#         doc = fitz.open(file_path)
#         for i, page in enumerate(doc):
#             text = page.get_text("text")
#             if text.strip():
#                 chunks.append({"content": text, "page": i+1, "source": file_path})

#     elif ext in [".docx", ".doc"]:
#         doc = Document(file_path)
#         for i, para in enumerate(doc.paragraphs):
#             if para.text.strip():
#                 chunks.append({"content": para.text, "page": i+1, "source": file_path})

#     elif ext in [".png", ".jpg", ".jpeg"]:
#         img = Image.open(file_path)
#         text = pytesseract.image_to_string(img)
#         chunks.append({"content": text, "page": 1, "source": file_path})

#     else:
#         raise ValueError("Unsupported file type")

#     return chunks



# def chunk_text(text, chunk_size=500, overlap=100):
#     chunks = []
#     start = 0

#     while start < len(text):
#         end = start + chunk_size
#         chunks.append(text[start:end])
#         start += chunk_size - overlap

#     return chunks



# def process_document(file_path):
#     base_name = os.path.splitext(os.path.basename(file_path))[0]

#     os.makedirs("storage", exist_ok=True)

#     embed_path = f"storage/{base_name}_embeddings.npy"
#     chunk_path = f"storage/{base_name}_chunks.pkl"
#     faiss_path = f"storage/{base_name}.index"

   
#     if os.path.exists(embed_path) and os.path.exists(faiss_path):
#         embeddings = np.load(embed_path)
#         index = faiss.read_index(faiss_path)

#         with open(chunk_path, "rb") as f:
#             final_chunks = pickle.load(f)

#         return final_chunks, embeddings, index

#     documents = parse_document(file_path)
#     final_chunks = []

#     for doc in documents:
#         small_chunks = chunk_text(doc["content"])
#         for chunk in small_chunks:
#             final_chunks.append({
#                 "content": chunk,
#                 "page": doc["page"],
#                 "source": doc["source"]
#             })

#     texts = [c["content"] for c in final_chunks]


#     embeddings = EMBED_MODEL.encode(texts, batch_size=32, normalize_embeddings=True)

#     dimension = embeddings.shape[1]
#     index = faiss.IndexFlatL2(dimension)
#     index.add(np.array(embeddings))

#     np.save(embed_path, embeddings)
#     faiss.write_index(index, faiss_path)

#     with open(chunk_path, "wb") as f:
#         pickle.dump(final_chunks, f)

#     return final_chunks, embeddings, index



# def retrieve_context(query, index, chunks, top_k=3):
#     query_embedding = EMBED_MODEL.encode([query], normalize_embeddings=True)

#     D, I = index.search(np.array(query_embedding), top_k)

#     retrieved = [chunks[i] for i in I[0]]

#     context = "\n\n".join([c["content"] for c in retrieved])
#     context = context[:1500]

#     return context, retrieved




















# # python -m uvicorn api:app --reload  command to run the API server
























import os
os.environ["TRANSFORMERS_OFFLINE"] = "1"
os.environ["HF_DATASETS_OFFLINE"] = "1"
import fitz
import pytesseract
from docx import Document
from PIL import Image
from sentence_transformers import SentenceTransformer
import numpy as np
import pickle
import faiss

pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

EMBED_MODEL = SentenceTransformer("all-MiniLM-L6-v2", device="cpu")
EMBED_MODEL.max_seq_length = 256


def parse_document(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    chunks = []

    if ext == ".pdf":
        doc = fitz.open(file_path)
        for i, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                chunks.append({"content": text, "page": i+1, "source": file_path})
            else:
               
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text = pytesseract.image_to_string(img)
                if text.strip():
                    chunks.append({"content": text, "page": i+1, "source": file_path})

    elif ext in [".docx", ".doc"]:
        doc = Document(file_path)
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                chunks.append({"content": para.text, "page": i+1, "source": file_path})

    elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp"]:
        img = Image.open(file_path)
    
        img = img.convert("L")  # grayscale
        img = img.point(lambda x: 0 if x < 140 else 255)  # threshold
        text = pytesseract.image_to_string(img)
        if text.strip():
            chunks.append({"content": text, "page": 1, "source": file_path})
        else:
            chunks.append({"content": "No readable text found in image.", "page": 1, "source": file_path})

    else:
        raise ValueError("Unsupported file type")

    return chunks



def chunk_text(text, chunk_size=500, overlap=100):
    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap

    return chunks

def process_document(file_path):
    base_name = os.path.splitext(os.path.basename(file_path))[0]

    os.makedirs("storage", exist_ok=True)

    embed_path = f"storage/{base_name}_embeddings.npy"
    chunk_path = f"storage/{base_name}_chunks.pkl"
    faiss_path = f"storage/{base_name}.index"

    if os.path.exists(embed_path) and os.path.exists(faiss_path):
        embeddings = np.load(embed_path)
        index = faiss.read_index(faiss_path)

        with open(chunk_path, "rb") as f:
            final_chunks = pickle.load(f)

        return final_chunks, embeddings, index

    documents = parse_document(file_path)
    final_chunks = []

    for doc in documents:
        small_chunks = chunk_text(doc["content"])
        for chunk in small_chunks:
            final_chunks.append({
                "content": chunk,
                "page": doc["page"],
                "source": doc["source"]
            })

    texts = [c["content"] for c in final_chunks]

    embeddings = EMBED_MODEL.encode(texts, batch_size=32, normalize_embeddings=True)

    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(np.array(embeddings))

    np.save(embed_path, embeddings)
    faiss.write_index(index, faiss_path)

    with open(chunk_path, "wb") as f:
        pickle.dump(final_chunks, f)

    return final_chunks, embeddings, index



def retrieve_context(query, index, chunks, top_k=3):
    query_embedding = EMBED_MODEL.encode([query], normalize_embeddings=True)

    D, I = index.search(np.array(query_embedding), top_k)

    retrieved = [chunks[i] for i in I[0]]

    context = "\n\n".join([c["content"] for c in retrieved])
    context = context[:800]

    return context, retrieved